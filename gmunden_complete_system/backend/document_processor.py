#!/usr/bin/env python3
"""
Document Processor für Gmunden Transparenz-System
Vollständige Dokumentenverarbeitung mit OCR
"""

import os
import logging
import hashlib
import mimetypes
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
import json
import re

# OCR und PDF-Verarbeitung
try:
    import pytesseract
    from PIL import Image
    import pdf2image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logging.warning("OCR-Module nicht verfügbar. OCR-Funktionen deaktiviert.")

# Office-Dokumente
try:
    from docx import Document as DocxDocument
    import openpyxl
    OFFICE_AVAILABLE = True
except ImportError:
    OFFICE_AVAILABLE = False
    logging.warning("Office-Module nicht verfügbar. Office-Dokument-Verarbeitung eingeschränkt.")

class DocumentProcessor:
    """Vollständige Dokumentenverarbeitung"""
    
    def __init__(self, upload_dir: str = "data/uploads", 
                 processed_dir: str = "data/processed"):
        self.upload_dir = Path(upload_dir)
        self.processed_dir = Path(processed_dir)
        
        # Verzeichnisse erstellen
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        self.processed_dir.mkdir(parents=True, exist_ok=True)
        
        # Unterstützte Dateitypen
        self.supported_types = {
            'pdf': ['.pdf'],
            'image': ['.jpg', '.jpeg', '.png', '.tiff', '.bmp'],
            'office': ['.docx', '.xlsx', '.pptx'],
            'text': ['.txt', '.csv', '.json', '.xml'],
            'archive': ['.zip', '.tar', '.gz']
        }
        
        # OCR-Konfiguration
        if OCR_AVAILABLE:
            self.ocr_config = r'--oem 3 --psm 6 -l deu+eng'
        
        logging.info("✅ DocumentProcessor initialisiert")
    
    def get_file_info(self, file_path: Path) -> Dict[str, Any]:
        """Extrahiere Datei-Informationen"""
        try:
            stat = file_path.stat()
            mime_type, _ = mimetypes.guess_type(str(file_path))
            
            # Datei-Hash berechnen
            with open(file_path, 'rb') as f:
                file_hash = hashlib.sha256(f.read()).hexdigest()
            
            return {
                'filename': file_path.name,
                'filepath': str(file_path),
                'size': stat.st_size,
                'size_mb': round(stat.st_size / (1024 * 1024), 2),
                'mime_type': mime_type,
                'extension': file_path.suffix.lower(),
                'created': datetime.fromtimestamp(stat.st_ctime),
                'modified': datetime.fromtimestamp(stat.st_mtime),
                'hash': file_hash
            }
        except Exception as e:
            logging.error(f"Fehler beim Extrahieren der Datei-Info: {e}")
            return {}
    
    def detect_file_type(self, file_path: Path) -> str:
        """Erkenne Dateityp"""
        extension = file_path.suffix.lower()
        
        for file_type, extensions in self.supported_types.items():
            if extension in extensions:
                return file_type
        
        return 'unknown'
    
    def extract_text_from_pdf(self, file_path: Path) -> Tuple[str, Dict]:
        """Extrahiere Text aus PDF mit OCR"""
        if not OCR_AVAILABLE:
            return "", {"error": "OCR nicht verfügbar"}
        
        try:
            # PDF zu Bildern konvertieren
            images = pdf2image.convert_from_path(str(file_path))
            
            extracted_text = ""
            ocr_stats = {
                'pages': len(images),
                'confidence_scores': [],
                'processing_time': 0
            }
            
            start_time = datetime.now()
            
            for i, image in enumerate(images):
                try:
                    # OCR auf jeder Seite
                    page_text = pytesseract.image_to_string(
                        image, 
                        config=self.ocr_config
                    )
                    
                    # Konfidenz-Score
                    try:
                        data = pytesseract.image_to_data(
                            image, 
                            config=self.ocr_config,
                            output_type=pytesseract.Output.DICT
                        )
                        confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
                        avg_confidence = sum(confidences) / len(confidences) if confidences else 0
                        ocr_stats['confidence_scores'].append(avg_confidence)
                    except:
                        ocr_stats['confidence_scores'].append(0)
                    
                    extracted_text += f"\n--- Seite {i+1} ---\n{page_text}\n"
                    
                except Exception as e:
                    logging.error(f"OCR-Fehler auf Seite {i+1}: {e}")
                    extracted_text += f"\n--- Seite {i+1} (Fehler) ---\n"
            
            ocr_stats['processing_time'] = (datetime.now() - start_time).total_seconds()
            ocr_stats['avg_confidence'] = sum(ocr_stats['confidence_scores']) / len(ocr_stats['confidence_scores']) if ocr_stats['confidence_scores'] else 0
            
            return extracted_text.strip(), ocr_stats
            
        except Exception as e:
            logging.error(f"PDF-OCR-Fehler: {e}")
            return "", {"error": str(e)}
    
    def extract_text_from_image(self, file_path: Path) -> Tuple[str, Dict]:
        """Extrahiere Text aus Bild mit OCR"""
        if not OCR_AVAILABLE:
            return "", {"error": "OCR nicht verfügbar"}
        
        try:
            image = Image.open(file_path)
            
            start_time = datetime.now()
            
            # OCR
            text = pytesseract.image_to_string(image, config=self.ocr_config)
            
            # Konfidenz-Score
            try:
                data = pytesseract.image_to_data(
                    image, 
                    config=self.ocr_config,
                    output_type=pytesseract.Output.DICT
                )
                confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
                avg_confidence = sum(confidences) / len(confidences) if confidences else 0
            except:
                avg_confidence = 0
            
            ocr_stats = {
                'processing_time': (datetime.now() - start_time).total_seconds(),
                'confidence': avg_confidence,
                'image_size': image.size
            }
            
            return text.strip(), ocr_stats
            
        except Exception as e:
            logging.error(f"Bild-OCR-Fehler: {e}")
            return "", {"error": str(e)}
    
    def extract_text_from_docx(self, file_path: Path) -> Tuple[str, Dict]:
        """Extrahiere Text aus Word-Dokument"""
        if not OFFICE_AVAILABLE:
            return "", {"error": "Office-Module nicht verfügbar"}
        
        try:
            doc = DocxDocument(file_path)
            
            text_parts = []
            stats = {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'sections': len(doc.sections)
            }
            
            # Paragraphen
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Tabellen
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts), stats
            
        except Exception as e:
            logging.error(f"DOCX-Fehler: {e}")
            return "", {"error": str(e)}
    
    def extract_text_from_xlsx(self, file_path: Path) -> Tuple[str, Dict]:
        """Extrahiere Text aus Excel-Datei"""
        if not OFFICE_AVAILABLE:
            return "", {"error": "Office-Module nicht verfügbar"}
        
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            text_parts = []
            stats = {
                'worksheets': len(workbook.worksheets),
                'total_cells': 0
            }
            
            for worksheet in workbook.worksheets:
                text_parts.append(f"\n--- {worksheet.title} ---")
                
                for row in worksheet.iter_rows(values_only=True):
                    row_text = []
                    for cell_value in row:
                        if cell_value is not None:
                            row_text.append(str(cell_value))
                            stats['total_cells'] += 1
                    
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts), stats
            
        except Exception as e:
            logging.error(f"XLSX-Fehler: {e}")
            return "", {"error": str(e)}
    
    def extract_text_from_txt(self, file_path: Path) -> Tuple[str, Dict]:
        """Extrahiere Text aus Textdatei"""
        try:
            # Encoding erkennen
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    
                    stats = {
                        'encoding': encoding,
                        'lines': len(text.split('\n')),
                        'characters': len(text)
                    }
                    
                    return text, stats
                    
                except UnicodeDecodeError:
                    continue
            
            return "", {"error": "Encoding nicht erkannt"}
            
        except Exception as e:
            logging.error(f"TXT-Fehler: {e}")
            return "", {"error": str(e)}
    
    def extract_metadata(self, file_path: Path, file_type: str) -> Dict[str, Any]:
        """Extrahiere Metadaten basierend auf Dateityp"""
        metadata = {}
        
        try:
            # Basis-Metadaten
            file_info = self.get_file_info(file_path)
            metadata.update(file_info)
            
            # Dateiname-Analyse
            filename = file_path.stem
            
            # Jahr aus Dateiname extrahieren
            year_match = re.search(r'20\d{2}', filename)
            if year_match:
                metadata['jahr'] = int(year_match.group())
            
            # Kategorie aus Dateiname ableiten
            filename_lower = filename.lower()
            if any(word in filename_lower for word in ['protokoll', 'sitzung', 'gemeinderat']):
                metadata['kategorie'] = 'protokolle'
            elif any(word in filename_lower for word in ['budget', 'finanz', 'rechnung', 'kosten']):
                metadata['kategorie'] = 'finanzen'
            elif any(word in filename_lower for word in ['bericht', 'statistik', 'analyse']):
                metadata['kategorie'] = 'berichte'
            else:
                metadata['kategorie'] = 'dokumente'
            
            # Tags aus Dateiname
            tags = []
            tag_patterns = {
                'gemeinderat': r'gemeinderat|gr|rat',
                'protokoll': r'protokoll|sitzung|meeting',
                'budget': r'budget|haushalt|finanz',
                'bericht': r'bericht|report|analyse',
                'statistik': r'statistik|zahlen|daten'
            }
            
            for tag, pattern in tag_patterns.items():
                if re.search(pattern, filename_lower):
                    tags.append(tag)
            
            metadata['tags'] = tags
            
            return metadata
            
        except Exception as e:
            logging.error(f"Metadaten-Fehler: {e}")
            return metadata
    
    def process_single_file(self, file_path: Path, 
                           enable_ocr: bool = True,
                           custom_metadata: Dict = None) -> Dict[str, Any]:
        """Verarbeite einzelne Datei vollständig"""
        try:
            # Datei-Info
            file_info = self.get_file_info(file_path)
            file_type = self.detect_file_type(file_path)
            
            # Basis-Ergebnis
            result = {
                'filename': file_info['filename'],
                'filepath': str(file_path),
                'file_type': file_type,
                'size_mb': file_info['size_mb'],
                'processed_at': datetime.now(),
                'success': False,
                'text_extracted': False,
                'ocr_used': False,
                'extracted_text': "",
                'processing_stats': {},
                'metadata': {},
                'errors': []
            }
            
            # Metadaten extrahieren
            metadata = self.extract_metadata(file_path, file_type)
            if custom_metadata:
                metadata.update(custom_metadata)
            result['metadata'] = metadata
            
            # Text-Extraktion basierend auf Dateityp
            extracted_text = ""
            processing_stats = {}
            
            if file_type == 'pdf' and enable_ocr:
                extracted_text, processing_stats = self.extract_text_from_pdf(file_path)
                result['ocr_used'] = True
                
            elif file_type == 'image' and enable_ocr:
                extracted_text, processing_stats = self.extract_text_from_image(file_path)
                result['ocr_used'] = True
                
            elif file_type == 'office':
                if file_path.suffix.lower() == '.docx':
                    extracted_text, processing_stats = self.extract_text_from_docx(file_path)
                elif file_path.suffix.lower() == '.xlsx':
                    extracted_text, processing_stats = self.extract_text_from_xlsx(file_path)
                    
            elif file_type == 'text':
                extracted_text, processing_stats = self.extract_text_from_txt(file_path)
            
            # Ergebnis aktualisieren
            if extracted_text:
                result['extracted_text'] = extracted_text
                result['text_extracted'] = True
                result['text_length'] = len(extracted_text)
                result['word_count'] = len(extracted_text.split())
            
            result['processing_stats'] = processing_stats
            
            # Fehler prüfen
            if 'error' in processing_stats:
                result['errors'].append(processing_stats['error'])
            else:
                result['success'] = True
            
            # Verarbeitete Datei speichern (optional)
            if result['success']:
                processed_file = self.processed_dir / f"{file_path.stem}_processed.json"
                with open(processed_file, 'w', encoding='utf-8') as f:
                    json.dump(result, f, ensure_ascii=False, indent=2, default=str)
            
            logging.info(f"✅ Datei verarbeitet: {file_path.name}")
            return result
            
        except Exception as e:
            logging.error(f"Verarbeitungsfehler für {file_path}: {e}")
            result['errors'].append(str(e))
            return result
    
    def process_bulk_files(self, file_paths: List[Path],
                          enable_ocr: bool = True,
                          custom_metadata: Dict = None,
                          chunk_size: int = 3,
                          max_retries: int = 3) -> Dict[str, Any]:
        """Verarbeite mehrere Dateien in Chunks"""
        
        total_files = len(file_paths)
        results = {
            'total_files': total_files,
            'processed': 0,
            'successful': 0,
            'failed': 0,
            'chunks_processed': 0,
            'processing_time': 0,
            'files': [],
            'errors': [],
            'stats': {
                'total_size_mb': 0,
                'total_text_length': 0,
                'ocr_files': 0,
                'file_types': {}
            }
        }
        
        start_time = datetime.now()
        
        # Dateien in Chunks aufteilen
        chunks = [file_paths[i:i + chunk_size] for i in range(0, total_files, chunk_size)]
        
        logging.info(f"🔄 Starte Bulk-Verarbeitung: {total_files} Dateien in {len(chunks)} Chunks")
        
        for chunk_idx, chunk in enumerate(chunks):
            logging.info(f"📦 Verarbeite Chunk {chunk_idx + 1}/{len(chunks)}: {len(chunk)} Dateien")
            
            for file_path in chunk:
                retry_count = 0
                file_result = None
                
                while retry_count < max_retries:
                    try:
                        file_result = self.process_single_file(
                            file_path, 
                            enable_ocr=enable_ocr,
                            custom_metadata=custom_metadata
                        )
                        
                        if file_result['success']:
                            break
                        else:
                            retry_count += 1
                            if retry_count < max_retries:
                                logging.warning(f"⚠️ Wiederhole {file_path.name} (Versuch {retry_count + 1})")
                    
                    except Exception as e:
                        retry_count += 1
                        if retry_count >= max_retries:
                            file_result = {
                                'filename': file_path.name,
                                'success': False,
                                'errors': [str(e)]
                            }
                
                # Ergebnisse sammeln
                if file_result:
                    results['files'].append(file_result)
                    results['processed'] += 1
                    
                    if file_result['success']:
                        results['successful'] += 1
                        
                        # Statistiken
                        if 'size_mb' in file_result:
                            results['stats']['total_size_mb'] += file_result['size_mb']
                        
                        if file_result['text_extracted']:
                            results['stats']['total_text_length'] += file_result.get('text_length', 0)
                        
                        if file_result['ocr_used']:
                            results['stats']['ocr_files'] += 1
                        
                        # Dateityp-Statistik
                        file_type = file_result.get('file_type', 'unknown')
                        results['stats']['file_types'][file_type] = results['stats']['file_types'].get(file_type, 0) + 1
                        
                    else:
                        results['failed'] += 1
                        results['errors'].extend(file_result.get('errors', []))
            
            results['chunks_processed'] += 1
        
        results['processing_time'] = (datetime.now() - start_time).total_seconds()
        
        logging.info(f"✅ Bulk-Verarbeitung abgeschlossen: {results['successful']}/{total_files} erfolgreich")
        
        return results
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """Hole Verarbeitungsstatistiken"""
        try:
            processed_files = list(self.processed_dir.glob("*_processed.json"))
            
            stats = {
                'total_processed': len(processed_files),
                'file_types': {},
                'total_size_mb': 0,
                'ocr_files': 0,
                'recent_files': []
            }
            
            for processed_file in processed_files[-10:]:  # Letzte 10
                try:
                    with open(processed_file, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                    
                    stats['recent_files'].append({
                        'filename': data.get('filename', ''),
                        'processed_at': data.get('processed_at', ''),
                        'success': data.get('success', False)
                    })
                    
                    if data.get('success'):
                        file_type = data.get('file_type', 'unknown')
                        stats['file_types'][file_type] = stats['file_types'].get(file_type, 0) + 1
                        stats['total_size_mb'] += data.get('size_mb', 0)
                        
                        if data.get('ocr_used'):
                            stats['ocr_files'] += 1
                
                except Exception as e:
                    logging.error(f"Fehler beim Lesen von {processed_file}: {e}")
            
            return stats
            
        except Exception as e:
            logging.error(f"Fehler beim Erstellen der Statistiken: {e}")
            return {}

# Globale Instanz
document_processor = DocumentProcessor()